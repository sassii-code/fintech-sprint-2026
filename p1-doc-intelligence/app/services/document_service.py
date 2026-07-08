import io
import pymupdf
import pdfplumber
import pytesseract
from PIL import Image, UnidentifiedImageError
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from fastapi import HTTPException, UploadFile


def extract_text_from_pdf(contents: bytes) -> str:
    text = ""
    try:
        doc = pymupdf.open(stream=contents, filetype="pdf")
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception:
        try:
            with pdfplumber.open(io.BytesIO(contents)) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
        except Exception:
            raise HTTPException(status_code=422, detail="Could not read PDF — file may be corrupted or malformed")
    return text


def extract_text_from_image(contents: bytes) -> str:
    try:
        image = Image.open(io.BytesIO(contents))
    except UnidentifiedImageError:
        raise HTTPException(status_code=422, detail="Could not read image — file may be corrupted or malformed")

    try:
        return pytesseract.image_to_string(image)
    except pytesseract.TesseractNotFoundError:
        raise HTTPException(status_code=503, detail="OCR engine is not available on this server")


def extract_text_from_docx(contents: bytes) -> str:
    try:
        doc = DocxDocument(io.BytesIO(contents))
    except PackageNotFoundError:
        raise HTTPException(status_code=422, detail="Could not read Word document — file may be corrupted or malformed")

    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text_from_txt(contents: bytes) -> str:
    try:
        return contents.decode("utf-8")
    except UnicodeDecodeError:
        return contents.decode("latin-1", errors="replace")


EXTENSION_HANDLERS = {
    ".pdf": extract_text_from_pdf,
    ".jpg": extract_text_from_image,
    ".jpeg": extract_text_from_image,
    ".png": extract_text_from_image,
    ".docx": extract_text_from_docx,
    ".txt": extract_text_from_txt,
}


async def file_to_text(file: UploadFile) -> str:
    if not file.filename or "." not in file.filename:
        supported = ", ".join(sorted(EXTENSION_HANDLERS))
        raise HTTPException(status_code=400, detail=f"Unsupported or missing file extension. Supported: {supported}")

    ext = "." + file.filename.rsplit(".", 1)[-1].lower()
    handler = EXTENSION_HANDLERS.get(ext)
    if not handler:
        supported = ", ".join(sorted(EXTENSION_HANDLERS))
        raise HTTPException(status_code=400, detail=f"Unsupported file type '{ext}'. Supported: {supported}")

    contents = await file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    text = handler(contents)
    if not text.strip():
        raise HTTPException(status_code=422, detail="Could not extract any text from the file")
    return text
